import os
from docx import Document 

def ingresar_datos():
    while True:
        nombre = input("Ingrese el nombre del trabajador: ")
        if nombre != "" and len(nombre) <= 30:
            break
        print("El nombre no puede estar vacío y debe tener máximo 30 caracteres.")
    
    while True:
        try:
            sueldo_base = int(input("Ingrese el sueldo base del trabajador: "))
            if sueldo_base >= 0:
                break
            print("El sueldo base debe ser un valor numérico positivo.")
        except ValueError:
            print("El sueldo base debe ser un valor numérico positivo.")
    
    while True:
        try:
            horas_extras = int(input("Ingrese el número de horas extras trabajadas: "))
            if horas_extras >= 0:
                break
            print("Las horas extras deben ser un valor numérico positivo.")
        except ValueError:
            print("Las horas extras deben ser un valor numérico positivo.")
    
    return nombre, sueldo_base, horas_extras

def calcular_liquidacion(sueldo_base, horas_extras):
    pago_horas_extras = horas_extras * (sueldo_base // 160) * 3 // 2
    total_ingresos = sueldo_base + pago_horas_extras
    descuento_fonasa = total_ingresos // 100 * 7
    descuento_afp = total_ingresos // 10
    sueldo_neto = total_ingresos - (descuento_fonasa + descuento_afp)
    return pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto

def mostrar_liquidacion(nombre, sueldo_base, pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto):
    print("\nLiquidación de sueldo de", nombre)
    print("Sueldo base: $", sueldo_base)
    print("Pago por horas extras: $", pago_horas_extras)
    print("Total de ingresos: $", total_ingresos)
    print("Descuento por FONASA: $", descuento_fonasa)
    print("Descuento por AFP: $", descuento_afp)
    print("Sueldo neto a pagar: $", sueldo_neto)

def generar_archivo_liquidacion(nombre, sueldo_base, pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto):
    archivo = Document()
    archivo.add_heading("Liquidación de sueldo", 0)
    archivo.add_paragraph("Nombre: " + nombre)
    archivo.add_paragraph("Sueldo base: $" + str(sueldo_base))
    archivo.add_paragraph("Pago por horas extras: $" + str(pago_horas_extras))
    archivo.add_paragraph("Total de ingresos: $" + str(total_ingresos))
    archivo.add_paragraph("Descuento por FONASA: $" + str(descuento_fonasa))
    archivo.add_paragraph("Descuento por AFP: $" + str(descuento_afp))
    archivo.add_paragraph("Sueldo neto a pagar: $" + str(sueldo_neto))
    
    nombre_archivo = "liquidacion_" + nombre + ".docx"
    archivo.save(nombre_archivo)
    print("Archivo de liquidación generado:", nombre_archivo)

def main():
    nombre, sueldo_base, horas_extras = ingresar_datos()
    pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto = calcular_liquidacion(sueldo_base, horas_extras)
    mostrar_liquidacion(nombre, sueldo_base, pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto)
    generar_archivo_liquidacion(nombre, sueldo_base, pago_horas_extras, total_ingresos, descuento_fonasa, descuento_afp, sueldo_neto)

if __name__ == "__main__":
    main()